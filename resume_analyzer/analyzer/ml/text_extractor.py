"""
Extracts raw text from uploaded resume files (PDF, DOCX, TXT).
"""
import os
import docx
import pdfplumber


class TextExtractionError(Exception):
    pass


def extract_text(file_path):
    """Dispatch extraction based on file extension. Returns plain text string."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext == ".docx":
        return _extract_from_docx(file_path)
    elif ext == ".txt":
        return _extract_from_txt(file_path)
    else:
        raise TextExtractionError(f"Unsupported file type: {ext}")


def _extract_from_pdf(file_path):
    text_chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
    except Exception as exc:
        raise TextExtractionError(f"Could not read PDF: {exc}")

    text = "\n".join(text_chunks).strip()
    if not text:
        raise TextExtractionError(
            "No selectable text found in this PDF. It may be a scanned image — "
            "please upload a text-based PDF or DOCX instead."
        )
    return text


def _extract_from_docx(file_path):
    try:
        document = docx.Document(file_path)
    except Exception as exc:
        raise TextExtractionError(f"Could not read DOCX: {exc}")

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text from tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    text = "\n".join(parts).strip()
    if not text:
        raise TextExtractionError("The DOCX file appears to be empty.")
    return text


def _extract_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    if not text:
        raise TextExtractionError("The TXT file appears to be empty.")
    return text
